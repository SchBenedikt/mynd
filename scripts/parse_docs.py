#!/usr/bin/env python3
"""
Document Parsing Pipeline
Converts PDF, DOCX, and Markdown files to clean Markdown for LightRAG ingestion.
Uses Docling for PDFs (best quality), python-docx for Word, preserves MD structure.
"""

import os
import sys
import logging
import argparse
from pathlib import Path
from typing import Optional, List
from dataclasses import dataclass

# PDF parsing
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Markdown parsing
import markdown
from markdown.extensions.toc import TocExtension
from markdown.extensions.tables import TableExtension
from markdown.extensions.fenced_code import FencedCodeExtension

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Result of document parsing."""
    content: str
    metadata: dict
    headings: List[dict]  # [{level, text, position}]
    tables: List[str]


def parse_pdf_docling(filepath: Path) -> ParsedDocument:
    """Parse PDF using Docling (IBM) - best quality for structured PDFs."""
    if not DOCLING_AVAILABLE:
        raise RuntimeError("Docling not installed. Run: pip install docling")
    
    logger.info(f"Parsing PDF with Docling: {filepath}")
    
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = True  # Enable OCR for scanned PDFs
    pipeline_options.do_table_structure = True
    pipeline_options.generate_page_images = False
    pipeline_options.generate_picture_images = False
    
    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: pipeline_options
        }
    )
    
    result = converter.convert(str(filepath))
    
    # Export to Markdown
    md_content = result.document.export_to_markdown()
    
    # Extract metadata
    metadata = {
        'source': str(filepath),
        'parser': 'docling',
        'pages': len(result.document.pages) if hasattr(result.document, 'pages') else 0,
    }
    
    # Extract headings from the document structure
    headings = []
    if hasattr(result.document, 'body') and hasattr(result.document.body, 'children'):
        for i, child in enumerate(result.document.body.children):
            if hasattr(child, 'label') and 'heading' in str(child.label).lower():
                level = getattr(child, 'level', 1)
                text = child.text if hasattr(child, 'text') else ''
                headings.append({'level': level, 'text': text, 'position': i})
    
    # Extract tables
    tables = []
    if hasattr(result.document, 'tables'):
        for table in result.document.tables:
            tables.append(table.export_to_markdown())
    
    return ParsedDocument(
        content=md_content,
        metadata=metadata,
        headings=headings,
        tables=tables,
    )


def parse_docx(filepath: Path) -> ParsedDocument:
    """Parse DOCX using python-docx, convert to Markdown."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    
    logger.info(f"Parsing DOCX: {filepath}")
    
    doc = DocxDocument(filepath)
    
    md_lines = []
    headings = []
    tables = []
    
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name.lower() if para.style else ''
        
        # Handle headings
        if style.startswith('heading'):
            level = int(style.replace('heading', '')) if style.replace('heading', '').isdigit() else 1
            text = para.text.strip()
            if text:
                md_lines.append(f"{'#' * level} {text}")
                headings.append({'level': level, 'text': text, 'position': len(md_lines)})
        else:
            text = para.text.strip()
            if text:
                # Bold/italic handling
                runs_text = []
                for run in para.runs:
                    t = run.text
                    if run.bold and run.italic:
                        t = f"***{t}***"
                    elif run.bold:
                        t = f"**{t}**"
                    elif run.italic:
                        t = f"*{t}*"
                    runs_text.append(t)
                md_lines.append(''.join(runs_text))
    
    # Extract tables
    for table in doc.tables:
        md_table = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            md_table.append('| ' + ' | '.join(cells) + ' |')
            if row_idx == 0:
                md_table.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        tables.append('\n'.join(md_table))
        md_lines.append('\n' + tables[-1] + '\n')
    
    metadata = {
        'source': str(filepath),
        'parser': 'python-docx',
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
    }
    
    return ParsedDocument(
        content='\n\n'.join(md_lines),
        metadata=metadata,
        headings=headings,
        tables=tables,
    )


def parse_markdown(filepath: Path) -> ParsedDocument:
    """Parse Markdown - preserve structure, extract headings."""
    logger.info(f"Parsing Markdown: {filepath}")
    
    content = filepath.read_text(encoding='utf-8')
    
    # Parse with extensions to extract TOC
    md = markdown.Markdown(extensions=[
        TocExtension(),
        TableExtension(),
        FencedCodeExtension(),
    ])
    md.convert(content)
    
    headings = []
    if hasattr(md, 'toc_tokens'):
        for token in md.toc_tokens:
            headings.append({
                'level': token['level'],
                'text': token['name'],
                'position': token.get('id', ''),
            })
    
    metadata = {
        'source': str(filepath),
        'parser': 'markdown',
        'size': len(content),
    }
    
    return ParsedDocument(
        content=content,
        metadata=metadata,
        headings=headings,
        tables=[],
    )


def parse_text(filepath: Path) -> ParsedDocument:
    """Parse plain text files."""
    logger.info(f"Parsing text file: {filepath}")
    
    content = filepath.read_text(encoding='utf-8')
    
    metadata = {
        'source': str(filepath),
        'parser': 'text',
        'size': len(content),
    }
    
    return ParsedDocument(
        content=content,
        metadata=metadata,
        headings=[],
        tables=[],
    )


def parse_document(filepath: Path, output_dir: Optional[Path] = None) -> Path:
    """
    Parse a document and save as Markdown.
    Returns path to the parsed Markdown file.
    """
    suffix = filepath.suffix.lower()
    
    if suffix == '.pdf':
        parsed = parse_pdf_docling(filepath)
    elif suffix == '.docx':
        parsed = parse_docx(filepath)
    elif suffix in ['.md', '.markdown']:
        parsed = parse_markdown(filepath)
    elif suffix in ['.txt', '.text']:
        parsed = parse_text(filepath)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
    
    # Determine output path
    if output_dir:
        rel_path = filepath.relative_to(filepath.anchor) if filepath.is_absolute() else filepath
        output_path = output_dir / rel_path.with_suffix('.md')
    else:
        output_path = filepath.with_suffix('.md')
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Add metadata header
    meta_lines = [
        '---',
        f'source: {parsed.metadata.get("source", "")}',
        f'parser: {parsed.metadata.get("parser", "")}',
        f'parsed_at: {__import__("datetime").datetime.now().isoformat()}',
        '---',
        '',
    ]
    
    full_content = '\n'.join(meta_lines) + parsed.content
    
    output_path.write_text(full_content, encoding='utf-8')
    logger.info(f"Saved parsed document: {output_path}")
    
    return output_path


def batch_parse(input_dir: Path, output_dir: Path, extensions: List[str] = None) -> List[Path]:
    """Parse all documents in a directory."""
    if extensions is None:
        extensions = ['.pdf', '.docx', '.md', '.markdown', '.txt', '.text']
    
    parsed_files = []
    
    for ext in extensions:
        for filepath in input_dir.rglob(f'*{ext}'):
            try:
                output_path = parse_document(filepath, output_dir)
                parsed_files.append(output_path)
            except Exception as e:
                logger.error(f"Failed to parse {filepath}: {e}")
    
    return parsed_files


def main():
    parser = argparse.ArgumentParser(description='Parse documents to Markdown')
    parser.add_argument('input', type=str, help='Input file or directory')
    parser.add_argument('-o', '--output', type=str, help='Output directory')
    parser.add_argument('--batch', action='store_true', help='Batch process directory')
    parser.add_argument('--extensions', type=str, default='.pdf,.docx,.md,.txt',
                        help='Comma-separated file extensions')
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    output_dir = Path(args.output) if args.output else None
    extensions = [e.strip() for e in args.extensions.split(',')]
    
    if args.batch or input_path.is_dir():
        if not output_dir:
            logger.error("Output directory required for batch mode")
            sys.exit(1)
        parsed = batch_parse(input_path, output_dir, extensions)
        print(f"Parsed {len(parsed)} files")
    else:
        if not input_path.exists():
            logger.error(f"File not found: {input_path}")
            sys.exit(1)
        output_path = parse_document(input_path, output_dir)
        print(f"Parsed: {output_path}")


if __name__ == '__main__':
    main()